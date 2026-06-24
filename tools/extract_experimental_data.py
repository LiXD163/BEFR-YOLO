"""Extract reported BE-FR YOLO experimental data from the manuscript.

This script organizes values already present in the paper. It never runs training,
validation, or inference, and it does not synthesize missing experimental results.
"""

from __future__ import annotations

import argparse
import csv
import json
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "results" / "organized_data"
MANUSCRIPT_SUFFIXES = {".docx", ".md", ".txt", ".pdf"}
TITLE_TERMS = (
    "BE-FR YOLO",
    "Progressive Blur Evolution",
    "Frequency Refinement",
    "Robust Vehicle Detection",
)

TABLE_SPECS = {
    2: (
        "ua_detrac_comparison",
        "table2_ua_detrac_comparison.csv",
        [
            "method",
            "clear_map50",
            "clear_map5095",
            "blur_light_map50",
            "blur_light_map5095",
            "blur_medium_map50",
            "blur_medium_map5095",
            "blur_heavy_map50",
            "blur_heavy_map5095",
        ],
    ),
    3: (
        "uavdt_comparison",
        "table3_uavdt_comparison.csv",
        [
            "method",
            "clear_map50",
            "clear_map5095",
            "blur_light_map50",
            "blur_light_map5095",
            "blur_medium_map50",
            "blur_medium_map5095",
            "blur_heavy_map50",
            "blur_heavy_map5095",
        ],
    ),
    4: (
        "bdd100k_realblur",
        "table4_bdd100k_realblur.csv",
        ["method", "precision", "recall", "map50", "map5095", "delta_map50_vs_yolov8"],
    ),
    5: (
        "ablation_ua_detrac",
        "table5_ablation_ua_detrac.csv",
        [
            "variant",
            "clear_map50",
            "clear_map5095",
            "blur_light_map50",
            "blur_light_map5095",
            "blur_medium_map50",
            "blur_medium_map5095",
            "blur_heavy_map50",
            "blur_heavy_map5095",
        ],
    ),
    6: (
        "statistical_significance",
        "table6_statistical_significance.csv",
        ["blur_level", "comparison", "mean_delta_map50", "test", "p_value", "significance"],
    ),
    7: (
        "sensitivity_lambda",
        "table7_sensitivity_lambda.csv",
        ["lambda", "clear_map50", "blur_light_map50", "blur_medium_map50", "blur_heavy_map50"],
    ),
    8: (
        "sensitivity_alpha_beta",
        "table8_sensitivity_alpha_beta.csv",
        ["alpha_beta", "clear_map50", "blur_light_map50", "blur_medium_map50", "blur_heavy_map50"],
    ),
    9: (
        "sensitivity_p_blur",
        "table9_sensitivity_p_blur.csv",
        ["p_range", "clear_map50", "blur_light_map50", "blur_medium_map50", "blur_heavy_map50"],
    ),
    10: (
        "sensitivity_l_range",
        "table10_sensitivity_l_range.csv",
        ["l_range", "clear_map50", "blur_light_map50", "blur_medium_map50", "blur_heavy_map50"],
    ),
    11: (
        "complexity_efficiency",
        "table11_complexity_efficiency.csv",
        ["method", "params_m", "flops_g", "fps", "latency_ms", "blur_heavy_map50"],
    ),
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract reported BE-FR YOLO experimental tables.")
    parser.add_argument("--paper", type=Path, default=None, help="Path to .docx, .md, or .txt manuscript.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--node", default=os.environ.get("BEFR_NODE") or shutil.which("node"))
    parser.add_argument("--skip-xlsx", action="store_true", help="Skip workbook creation.")
    return parser.parse_args()


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u00a0", " ")).strip()


def title_score(text: str) -> int:
    lowered = text.lower()
    return sum(term.lower() in lowered for term in TITLE_TERMS)


def candidate_title_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document

            document = Document(path)
            return "\n".join(paragraph.text for paragraph in document.paragraphs[:12])
        except Exception:
            return ""
    if path.suffix.lower() in {".md", ".txt"}:
        try:
            return path.read_text(encoding="utf-8")[:4000]
        except Exception:
            return ""
    return ""


def find_manuscript(explicit: Path | None) -> Path:
    if explicit is not None:
        path = explicit.expanduser().resolve()
        if not path.is_file():
            raise FileNotFoundError(f"Manuscript not found: {path}")
        return path
    candidates = [
        path
        for path in Path.cwd().rglob("*")
        if path.is_file()
        and path.suffix.lower() in MANUSCRIPT_SUFFIXES
        and not path.name.lower().startswith("readme")
        and "results" not in path.parts
    ]
    if not candidates:
        raise FileNotFoundError(
            "No manuscript (.docx, .pdf, .md, or .txt) was found in the current directory. "
            "Pass an explicit path with --paper."
        )
    scored = sorted(
        ((title_score(path.name + "\n" + candidate_title_text(path)), path) for path in candidates),
        key=lambda item: (-item[0], str(item[1])),
    )
    if scored[0][0] == 0:
        likely = [path for path in candidates if re.search(r"jvcir|be-fr|manuscript|paper", path.name, re.IGNORECASE)]
        if not likely:
            raise FileNotFoundError(
                "Files with manuscript-like extensions were found, but none matched the BE-FR YOLO title. "
                "Pass the intended file with --paper."
            )
        return sorted(likely)[0].resolve()
    return scored[0][1].resolve()


def direct_docx_rows(table: Any) -> list[list[str]]:
    """Read physical XML cells so grid-spanned cells are not duplicated by python-docx."""
    from docx.oxml.ns import qn

    rows: list[list[str]] = []
    for tr in table._tbl.tr_lst:
        row = []
        for tc in tr.tc_lst:
            text = "".join(node.text or "" for node in tc.iter() if node.tag == qn("w:t"))
            row.append(normalize_text(text))
        rows.append(row)
    return rows


def read_docx(path: Path) -> tuple[str, dict[int, list[list[str]]], list[str]]:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError("DOCX extraction requires python-docx. Install it with `pip install python-docx`.") from exc

    document = Document(path)
    paragraphs = [normalize_text(paragraph.text) for paragraph in document.paragraphs if normalize_text(paragraph.text)]
    full_text = "\n".join(paragraphs)
    table_map: dict[int, list[list[str]]] = {}
    pending_number: int | None = None

    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            text = normalize_text(paragraph.text)
            match = re.match(r"Table\s+(\d+)\.", text, flags=re.IGNORECASE)
            if match:
                pending_number = int(match.group(1))
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            if pending_number is not None:
                table_map[pending_number] = direct_docx_rows(table)
                pending_number = None
    return full_text, table_map, paragraphs


def parse_markdown_tables(text: str) -> dict[int, list[list[str]]]:
    lines = text.splitlines()
    tables: dict[int, list[list[str]]] = {}
    index = 0
    while index < len(lines):
        match = re.match(r"\s*Table\s+(\d+)\.", lines[index], flags=re.IGNORECASE)
        if not match:
            index += 1
            continue
        number = int(match.group(1))
        index += 1
        while index < len(lines) and "|" not in lines[index]:
            index += 1
        rows = []
        while index < len(lines) and "|" in lines[index]:
            cells = [normalize_text(cell) for cell in lines[index].strip().strip("|").split("|")]
            if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                rows.append(cells)
            index += 1
        if rows:
            tables[number] = rows
    return tables


def read_manuscript(path: Path) -> tuple[str, dict[int, list[list[str]]], list[str]]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx(path)
    if suffix in {".md", ".txt"}:
        text = path.read_text(encoding="utf-8")
        return text, parse_markdown_tables(text), [normalize_text(line) for line in text.splitlines() if normalize_text(line)]
    if suffix == ".pdf":
        raise RuntimeError(
            "PDF table extraction is not enabled because merged scientific tables are unreliable. "
            "Convert the manuscript to .docx or .txt and rerun."
        )
    raise RuntimeError(f"Unsupported manuscript type: {suffix}")


def first_match(pattern: str, text: str, cast: Any = str, flags: int = re.IGNORECASE | re.DOTALL) -> Any:
    match = re.search(pattern, text, flags)
    if not match:
        return None
    value = match.group(1)
    if cast in {float, int}:
        value = value.rstrip(".,;:")
    return cast(value)


def parse_scientific_number(value: str | None) -> float | None:
    if value is None:
        return None
    superscript = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹⁻", "0123456789-")
    cleaned = value.translate(superscript).replace("×", "e").replace(" ", "").rstrip(".,;:")
    cleaned = re.sub(r"e10", "e", cleaned)
    try:
        return float(cleaned)
    except ValueError:
        return None


def list_match(pattern: str, text: str, cast: Any = float) -> list[Any] | None:
    match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
    if not match:
        return None
    return [cast(match.group(1)), cast(match.group(2))]


def extract_config(text: str) -> tuple[dict[str, Any], list[str]]:
    gpu = first_match(r"on an\s+(.+?GPU)", text)
    if gpu and gpu.endswith(" GPU"):
        gpu = gpu[:-4]
    config = {
        "model": {
            "base_detector": first_match(r"built upon\s+([A-Za-z0-9-]+)", text),
            "input_size": first_match(r"resized to\s+(\d+)\s*[×x]\s*\d+", text, int),
            "epochs": first_match(r"trained for\s+(\d+)\s+epochs", text, int),
            "batch_size": first_match(r"batch size of\s+(\d+)", text, int),
            "optimizer": first_match(r"optimizer was\s+([A-Za-z0-9-]+)", text),
            "initial_learning_rate": first_match(r"initial learning rate of\s+([0-9.]+)", text, float),
            "momentum": first_match(r"momentum of\s+([0-9.]+)", text, float),
            "weight_decay": parse_scientific_number(first_match(r"weight decay of\s+([0-9×x10⁰¹²³⁴⁵⁶⁷⁸⁹⁻.^-]+)", text)),
            "gpu": gpu,
        },
        "be_net": {
            "blur_kernel_length_range": list_match(r"blur kernel length was constrained within\s*\[(\d+)\s*,\s*(\d+)\]", text, int),
            "blur_injection_probability_range": list_match(
                r"blur injection probability was constrained within\s*\[([0-9.]+)\s*,\s*([0-9.]+)\]", text
            ),
            "alpha": first_match(r"α\s*=\s*([0-9.]+)", text, float),
            "beta": first_match(r"β\s*=\s*([0-9.]+)", text, float),
            "pgn_hidden_dim": first_match(r"hidden dimension of\s+(\d+)", text, int),
            "theta_range": list_match(
                r"motion direction was randomly sampled from\s*\[([0-9.]+)°?\s*,\s*([0-9.]+)°?\)",
                text,
                lambda value: int(float(value)),
            ),
            "gaussian_noise_sigma": first_match(r"σn is set to\s+([0-9.]+)", text, float),
        },
        "fr_net": {
            "insertion_levels": ["P3", "P4", "P5"] if re.search(r"at P3, P4, and P5", text) else None,
            "spectral_consistency_loss_weight": first_match(r"spectral consistency loss weight was set to\s+λ\s*=\s*([0-9.]+)", text, float),
            "spatial_branch": "identity" if re.search(r"spatial branch was implemented as an identity mapping", text) else None,
            "fusion": "residual_addition" if re.search(r"fused with the original feature through residual addition", text) else None,
        },
        "synthetic_blur_protocol": {
            "clear": "original_images" if re.search(r"original testing images without synthetic blur were used as the Clear setting", text) else None,
            "blur_light": [9, 11, 13] if re.search(r"Blur-Light with L\s*∈\s*\{9\s*,\s*11\s*,\s*13\}", text) else None,
            "blur_medium": [15, 17, 19] if re.search(r"Blur-Medium with L\s*∈\s*\{15\s*,\s*17\s*,\s*19\}", text) else None,
            "blur_heavy": [21, 23, 25] if re.search(r"Blur-Heavy with L\s*∈\s*\{21\s*,\s*23\s*,\s*25\}", text) else None,
        },
        "bdd100k_realblur_subset": {
            "source": "BDD100K validation set" if re.search(r"BDD100K validation set", text) else None,
            "candidate_complex_images": first_match(r"([\d,]+)\s+complex traffic images were initially retained", text, lambda x: int(x.replace(",", ""))),
            "motion_blur_images_after_screening": first_match(r"and\s+([\d,]+)\s+images with visible motion-induced", text, lambda x: int(x.replace(",", ""))),
            "final_retained_images": first_match(r"final real motion-blur validation subset contains\s+([\d,]+)\s+images", text, lambda x: int(x.replace(",", ""))),
            "merged_categories": ["car", "bus", "truck", "motor"] if re.search(r"car, bus, truck, and motor", text) else None,
            "target_class": "vehicle" if re.search(r"merged into the vehicle class", text) else None,
        },
    }
    missing: list[str] = []

    def walk(node: Any, prefix: str = "") -> None:
        if isinstance(node, dict):
            for key, value in node.items():
                walk(value, f"{prefix}.{key}" if prefix else key)
        elif node is None:
            missing.append(prefix)

    walk(config)
    return config, missing


def normalize_method(value: str) -> str:
    value = normalize_text(value)
    value = re.sub(r"\s*-\s*", "-", value)
    return value.replace("BE-FR", "BE-FR")


def parse_float(value: str) -> float | None:
    value = normalize_text(value)
    if value in {"", "\\", "-", "—", "–"}:
        return None
    cleaned = re.sub(r"\s+", "", value)
    try:
        return float(cleaned)
    except ValueError as exc:
        raise ValueError(f"Expected numeric table cell, received {value!r}") from exc


def parse_range(value: str) -> str:
    return re.sub(r"\s+", "", normalize_text(value))


def parse_table(number: int, rows: list[list[str]]) -> list[dict[str, Any]]:
    if number not in TABLE_SPECS:
        raise KeyError(f"Unsupported requested table number: {number}")
    columns = TABLE_SPECS[number][2]
    if number in {2, 3, 5}:
        data_rows = rows[2:]
    else:
        data_rows = rows[1:]

    records: list[dict[str, Any]] = []
    for row in data_rows:
        if not row or not normalize_text(row[0]):
            continue
        if len(row) != len(columns):
            raise ValueError(f"Table {number} row has {len(row)} physical cells; expected {len(columns)}: {row}")
        record: dict[str, Any] = {}
        for index, (column, value) in enumerate(zip(columns, row)):
            if index == 0:
                if column in {"method", "variant"}:
                    record[column] = normalize_method(value)
                elif column in {"p_range", "l_range"}:
                    record[column] = parse_range(value)
                else:
                    record[column] = normalize_text(value)
            elif column in {"comparison", "test", "significance", "blur_level"}:
                record[column] = normalize_text(value)
            else:
                record[column] = parse_float(value)
        records.append(record)
    return records


def write_csv(path: Path, rows: list[dict[str, Any]], columns: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        raise ValueError(f"Refusing to write empty CSV: {path}")
    fieldnames = columns or list(rows[0])
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def yaml_scalar(value: Any) -> str:
    if value is None:
        return "null"
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, list):
        return "[" + ", ".join(yaml_scalar(item) for item in value) + "]"
    text = str(value)
    if re.fullmatch(r"[A-Za-z0-9_.-]+", text):
        return text
    return json.dumps(text, ensure_ascii=False)


def dump_yaml(data: dict[str, Any], indent: int = 0) -> str:
    lines: list[str] = []
    for key, value in data.items():
        prefix = " " * indent + f"{key}:"
        if isinstance(value, dict):
            lines.append(prefix)
            lines.append(dump_yaml(value, indent + 2).rstrip())
        else:
            lines.append(f"{prefix} {yaml_scalar(value)}")
    return "\n".join(lines) + "\n"


def method_record(rows: list[dict[str, Any]], method: str, field: str = "method") -> dict[str, Any]:
    for row in rows:
        if row[field] == method:
            return row
    raise KeyError(f"Missing {field}={method}")


def derived_summary(tables: dict[str, list[dict[str, Any]]]) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []

    def add(category: str, metric: str, baseline: float, proposed: float) -> None:
        output.append(
            {
                "category": category,
                "metric": metric,
                "baseline": baseline,
                "proposed_or_comparison": proposed,
                "value": round(proposed - baseline, 6),
            }
        )

    for key, category in (
        ("ua_detrac_comparison", "UA-DETRAC Blur-Heavy"),
        ("uavdt_comparison", "UAVDT Blur-Heavy"),
    ):
        baseline = method_record(tables[key], "YOLOv8")
        proposed = method_record(tables[key], "BE-FR YOLO")
        add(category, "mAP50 improvement", baseline["blur_heavy_map50"], proposed["blur_heavy_map50"])
        add(category, "mAP50-95 improvement", baseline["blur_heavy_map5095"], proposed["blur_heavy_map5095"])

    bdd = tables["bdd100k_realblur"]
    add("BDD100K real-blur", "mAP50 improvement", method_record(bdd, "YOLOv8")["map50"], method_record(bdd, "BE-FR YOLO")["map50"])
    add("BDD100K real-blur", "mAP50-95 improvement", method_record(bdd, "YOLOv8")["map5095"], method_record(bdd, "BE-FR YOLO")["map5095"])

    ablation = tables["ablation_ua_detrac"]
    proposed = method_record(ablation, "BE-FR YOLO", "variant")["blur_heavy_map50"]
    for variant in ("BE-only", "FR-only", "YOLOv8"):
        baseline = method_record(ablation, variant, "variant")["blur_heavy_map50"]
        add("Ablation Blur-Heavy", f"BE-FR YOLO minus {variant} mAP50", baseline, proposed)

    complexity = tables["complexity_efficiency"]
    baseline = method_record(complexity, "YOLOv8")
    proposed = method_record(complexity, "BE-FR YOLO")
    for metric, field in (
        ("parameter change (M)", "params_m"),
        ("FLOPs change (G)", "flops_g"),
        ("FPS change", "fps"),
        ("latency change (ms)", "latency_ms"),
    ):
        add("Complexity", metric, baseline[field], proposed[field])
    return output


def validation_report(tables: dict[str, list[dict[str, Any]]], source: Path) -> str:
    checks: list[tuple[str, float, float]] = []

    def add(label: str, actual: float, claimed: float) -> None:
        checks.append((label, actual, claimed))

    ua = tables["ua_detrac_comparison"]
    ua_base, ua_proposed = method_record(ua, "YOLOv8"), method_record(ua, "BE-FR YOLO")
    add("UA-DETRAC Blur-Heavy YOLOv8 mAP50", ua_base["blur_heavy_map50"], 0.714)
    add("UA-DETRAC Blur-Heavy BE-FR YOLO mAP50", ua_proposed["blur_heavy_map50"], 0.818)
    add("UA-DETRAC Blur-Heavy mAP50 improvement", ua_proposed["blur_heavy_map50"] - ua_base["blur_heavy_map50"], 0.104)

    uav = tables["uavdt_comparison"]
    uav_base, uav_proposed = method_record(uav, "YOLOv8"), method_record(uav, "BE-FR YOLO")
    add("UAVDT Blur-Heavy YOLOv8 mAP50", uav_base["blur_heavy_map50"], 0.171)
    add("UAVDT Blur-Heavy BE-FR YOLO mAP50", uav_proposed["blur_heavy_map50"], 0.286)
    add("UAVDT Blur-Heavy mAP50 improvement", uav_proposed["blur_heavy_map50"] - uav_base["blur_heavy_map50"], 0.115)

    bdd = tables["bdd100k_realblur"]
    bdd_base, bdd_proposed = method_record(bdd, "YOLOv8"), method_record(bdd, "BE-FR YOLO")
    add("BDD100K real-blur YOLOv8 mAP50", bdd_base["map50"], 0.446)
    add("BDD100K real-blur BE-FR YOLO mAP50", bdd_proposed["map50"], 0.556)
    add("BDD100K real-blur mAP50 improvement", bdd_proposed["map50"] - bdd_base["map50"], 0.110)

    complexity = tables["complexity_efficiency"]
    base, proposed = method_record(complexity, "YOLOv8"), method_record(complexity, "BE-FR YOLO")
    for label, actual, claimed in (
        ("YOLOv8 parameters (M)", base["params_m"], 3.2),
        ("BE-FR YOLO parameters (M)", proposed["params_m"], 4.4),
        ("Parameter increase (M)", proposed["params_m"] - base["params_m"], 1.2),
        ("YOLOv8 FLOPs (G)", base["flops_g"], 8.7),
        ("BE-FR YOLO FLOPs (G)", proposed["flops_g"], 11.6),
        ("FLOPs increase (G)", proposed["flops_g"] - base["flops_g"], 2.9),
        ("YOLOv8 FPS", base["fps"], 112.4),
        ("BE-FR YOLO FPS", proposed["fps"], 109.7),
    ):
        add(label, actual, claimed)

    lines = [
        "# Experimental Data Validation Report",
        "",
        f"- Manuscript: `{source}`",
        "- Scope: consistency between requested textual claims and extracted manuscript tables.",
        "- Status rule: PASS for exact agreement within 0.0005; WARNING within 0.0015; otherwise FAIL.",
        "",
        "| Status | Claim | Table value | Claimed value | Difference |",
        "|---|---|---:|---:|---:|",
    ]
    counts = {"PASS": 0, "WARNING": 0, "FAIL": 0}
    for label, actual, claimed in checks:
        difference = actual - claimed
        absolute = abs(difference)
        status = "PASS" if absolute <= 0.0005 else "WARNING" if absolute <= 0.0015 else "FAIL"
        counts[status] += 1
        lines.append(f"| {status} | {label} | {actual:.6f} | {claimed:.6f} | {difference:.6f} |")
    lines.extend(
        [
            "",
            "## Summary",
            "",
            f"- PASS: {counts['PASS']}",
            f"- WARNING: {counts['WARNING']}",
            f"- FAIL: {counts['FAIL']}",
            "",
            "## Extraction Notes",
            "",
            "- UAVDT Table 3 contains visually merged cells with three-column grid spans. The extractor reads physical DOCX XML cells, yielding the intended method plus eight metric values.",
            "- Statistical significance values are sequence-level mean differences and are preserved separately from differences computed from aggregate Table 5 values.",
        ]
    )
    return "\n".join(lines) + "\n"


def long_map50(rows: list[dict[str, Any]], dataset: str, name_field: str = "method") -> list[dict[str, Any]]:
    output = []
    for row in rows:
        for level, field in (
            ("Clear", "clear_map50"),
            ("Blur-Light", "blur_light_map50"),
            ("Blur-Medium", "blur_medium_map50"),
            ("Blur-Heavy", "blur_heavy_map50"),
        ):
            output.append({"dataset": dataset, "method": row[name_field], "blur_level": level, "metric": "mAP50", "value": row[field]})
    return output


def flatten_config(config: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []

    def walk(node: Any, prefix: str = "") -> None:
        if isinstance(node, dict):
            for key, value in node.items():
                walk(value, f"{prefix}.{key}" if prefix else key)
        else:
            rows.append({"field": prefix, "value": json.dumps(node, ensure_ascii=False) if isinstance(node, list) else node})

    walk(config)
    return rows


def workbook_payload(config: dict[str, Any], tables: dict[str, list[dict[str, Any]]]) -> dict[str, Any]:
    sheet_names = {
        "ua_detrac_comparison": "Table2_UA_DETRAC",
        "uavdt_comparison": "Table3_UAVDT",
        "bdd100k_realblur": "Table4_BDD100K",
        "ablation_ua_detrac": "Table5_Ablation",
        "statistical_significance": "Table6_Significance",
        "sensitivity_lambda": "Table7_Lambda",
        "sensitivity_alpha_beta": "Table8_AlphaBeta",
        "sensitivity_p_blur": "Table9_PBlur",
        "sensitivity_l_range": "Table10_LRange",
        "complexity_efficiency": "Table11_Complexity",
    }
    sheets = [{"name": "Config", "rows": flatten_config(config)}]
    sheets.extend({"name": sheet_names[key], "rows": tables[key]} for key in sheet_names)
    return {"sheets": sheets}


def create_xlsx(output_dir: Path, payload: dict[str, Any], node: str | None) -> None:
    if not node:
        raise RuntimeError("XLSX creation requires Node.js and @oai/artifact-tool. Pass --node or set BEFR_NODE.")
    payload_path = output_dir / ".workbook_payload.json"
    payload_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    builder = ROOT / "tools" / "build_experimental_workbook.mjs"
    destination = output_dir / "befr_experimental_results.xlsx"
    environment = os.environ.copy()
    adjacent_modules = Path(node).resolve().parent.parent / "node_modules"
    if adjacent_modules.is_dir() and not environment.get("NODE_PATH"):
        environment["NODE_PATH"] = str(adjacent_modules)
    subprocess.run([node, str(builder), str(payload_path), str(destination)], check=True, cwd=ROOT, env=environment)
    payload_path.unlink(missing_ok=True)


def write_missing_report(path: Path, missing: list[str], source: Path) -> None:
    lines = ["# Missing Field Report", "", f"- Manuscript: `{source}`", ""]
    if missing:
        lines.append("The following requested fields were not found and were written as `null`:")
        lines.extend(f"- `{field}`" for field in missing)
    else:
        lines.append("All requested experiment-configuration fields were found in the manuscript.")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    args = parse_args()
    paper = find_manuscript(args.paper)
    print(f"Using manuscript: {paper}")
    text, raw_tables, _paragraphs = read_manuscript(paper)
    config, missing = extract_config(text)

    absent_tables = [number for number in TABLE_SPECS if number not in raw_tables]
    if absent_tables:
        raise RuntimeError(f"Required manuscript tables not found: {absent_tables}")
    tables = {TABLE_SPECS[number][0]: parse_table(number, raw_tables[number]) for number in TABLE_SPECS}

    output_dir = args.output.resolve()
    tables_dir = output_dir / "tables"
    plot_dir = output_dir / "plot_data"
    output_dir.mkdir(parents=True, exist_ok=True)
    tables_dir.mkdir(parents=True, exist_ok=True)
    plot_dir.mkdir(parents=True, exist_ok=True)

    (output_dir / "experiment_config.yaml").write_text(dump_yaml(config), encoding="utf-8")
    for number, (key, filename, columns) in TABLE_SPECS.items():
        write_csv(tables_dir / filename, tables[key], columns)

    unified = {"experiment_config": config, "tables": tables}
    (output_dir / "befr_experimental_results.json").write_text(
        json.dumps(unified, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    write_csv(
        output_dir / "derived_summary.csv",
        derived_summary(tables),
        ["category", "metric", "baseline", "proposed_or_comparison", "value"],
    )
    (output_dir / "data_validation_report.md").write_text(validation_report(tables, paper), encoding="utf-8")
    write_missing_report(output_dir / "missing_fields_report.md", missing, paper)

    write_csv(plot_dir / "plot_main_comparison_ua_detrac_map50.csv", long_map50(tables["ua_detrac_comparison"], "UA-DETRAC"))
    write_csv(plot_dir / "plot_main_comparison_uavdt_map50.csv", long_map50(tables["uavdt_comparison"], "UAVDT"))
    write_csv(plot_dir / "plot_ablation_map50.csv", long_map50(tables["ablation_ua_detrac"], "UA-DETRAC", "variant"))
    for key, filename in (
        ("sensitivity_lambda", "plot_sensitivity_lambda.csv"),
        ("sensitivity_alpha_beta", "plot_sensitivity_alpha_beta.csv"),
        ("sensitivity_p_blur", "plot_sensitivity_p_blur.csv"),
        ("sensitivity_l_range", "plot_sensitivity_l_range.csv"),
        ("complexity_efficiency", "plot_complexity_accuracy_tradeoff.csv"),
    ):
        write_csv(plot_dir / filename, tables[key])

    if not args.skip_xlsx:
        create_xlsx(output_dir, workbook_payload(config, tables), args.node)

    print("Created experimental data outputs:")
    for path in sorted(path for path in output_dir.rglob("*") if path.is_file()):
        print(f"  {path}")
    print("Extracted table rows:")
    for key, rows in tables.items():
        print(f"  {key}: {len(rows)}")
    print(f"Validation report: {output_dir / 'data_validation_report.md'}")
    print("Missing fields: " + (", ".join(missing) if missing else "none"))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(1) from None
